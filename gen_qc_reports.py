#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""按科室分别生成 Word 质量控制数据报表（含雷达图），并打包为 ZIP。
排名在全院有核心临床数据（出院人次非空）的科室间计算。
"""
import csv, re, sys, zipfile, argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

sys.stdout.reconfigure(encoding='utf-8')
# 字体配置（Windows 优先 Microsoft YaHei）
plt.rcParams['font.sans-serif'] = ['Microsoft YaHei', 'SimHei', 'STHeiti', 'SimSun', 'Arial Unicode MS']
plt.rcParams['axes.unicode_minus'] = False
from matplotlib.font_manager import FontProperties
_SYSTEM_FONT = FontProperties(family='Microsoft YaHei')

_HERE = Path(__file__).resolve().parent
_DATA_DIR = _HERE / "data"  # 默认数据目录（相对脚本位置）

_ap = argparse.ArgumentParser(description="分科室质量控制报表生成器")
_ap.add_argument("--data", default=str(_DATA_DIR / "质控数据提取结果.csv"))
_ap.add_argument("--ddds", default=str(_DATA_DIR / "ddds目标值.csv"))
_ap.add_argument("--prev", default="", help="上月质控数据 CSV（可选）")
_ap.add_argument("--out", default=str(_DATA_DIR / "科室单报表"))
_ap.add_argument("--zip", default=str(_DATA_DIR / "科室质控报表.zip"))
_a = _ap.parse_args()

BASE = Path(_a.out).parent
DATA_CSV = Path(_a.data)
DDDS_TARGET_CSV = Path(_a.ddds)
OUT_DIR = Path(_a.out)
OUT_DIR.mkdir(parents=True, exist_ok=True)
ZIP_PATH = Path(_a.zip)
RADAR_DIR = OUT_DIR / "radar"
RADAR_DIR.mkdir(parents=True, exist_ok=True)

for _p in (DATA_CSV, DDDS_TARGET_CSV):
    if not _p.exists():
        sys.exit(f"错误：找不到数据文件 {_p}。请用 --data/--ddds 指定路径，或将数据放入 {_DATA_DIR}。")

CLR_TITLE = RGBColor(0x1F, 0x3A, 0x5F)
CLR_HEADER_BG = "1F3A5F"
CLR_BLACK = RGBColor(0x33, 0x33, 0x33)
CLR_RED = RGBColor(0xE5, 0x39, 0x35)
CLR_GREEN = RGBColor(0x00, 0xB0, 0x50)
CLR_BLUE = RGBColor(0x1F, 0x75, 0xFE)
CLR_ORANGE = RGBColor(0xFF, 0x99, 0x33)
CLR_GRAY = RGBColor(0x99, 0x99, 0x99)
QUARTILE_COLOR = {"A": CLR_GREEN, "B": CLR_BLUE, "C": CLR_ORANGE, "D": CLR_RED}
QUARTILE_RADAR_COLOR = {"A": "#00B050", "B": "#1F75FE", "C": "#FF9933", "D": "#E53935"}

# 指标分类
GROUPS = {
    "能力及效率": [
        "出院人次", "门（急）诊人次", "医疗服务收入占比", "门诊次均费用", "门诊次均药品费用",
        "住院次均费用", "住院患者均次药品费用", "CMI值", "平均住院日", "床位使用率",
        "出院患者手术占比", "出院患者四级手术占比", "出院患者微创手术占比",
    ],
    "医疗质量与安全": [
        "疑难病例讨论数", "死亡例数（住院）", "(48小时)再入院例数", "再次手术例数",
        "住院超30天病例数", "非医嘱离院例数（住院）", "医疗不良事件上报数",
        "检验危急值处理率", "临床路径入径率", "出院病历3个工作日归档率",
        "普通会诊未及时完成数", "督查扣分", "DDDs",
    ],
}
ALL_METRICS = [m for g in GROUPS.values() for m in g]

# 科室名称别名（上月→本月映射）
DEPT_ALIASES = {
    'ICU': '重症医学科',
    '康复科': '康复医学科',
    '整形外科': '烧伤整形科',
    '精神心理科': '临床心理科',
    '烧伤整形科': '烧伤科',
}

POSITIVE = {
    "出院人次", "门（急）诊人次", "CMI值", "床位使用率", "医疗服务收入占比",
    "检验危急值处理率", "疑难病例讨论数", "出院病历3个工作日归档率", "临床路径入径率",
    "出院患者手术占比", "出院患者四级手术占比", "出院患者微创手术占比",
}
NEGATIVE = {
    "门诊次均费用", "住院次均费用", "平均住院日", "住院患者均次药品费用", "门诊次均药品费用",
    "死亡例数（住院）", "(48小时)再入院例数", "住院超30天病例数", "非医嘱离院例数（住院）",
    "普通会诊未及时完成数", "督查扣分", "再次手术例数",
}
COUNT_METRICS = {
    "出院人次", "门（急）诊人次", "死亡例数（住院）", "(48小时)再入院例数",
    "住院超30天病例数", "非医嘱离院例数（住院）", "疑难病例讨论数",
    "普通会诊未及时完成数", "医疗不良事件上报数", "再次手术例数",
}
PERCENT_METRICS = {
    "医疗服务收入占比", "床位使用率", "出院患者手术占比", "出院患者四级手术占比",
    "出院患者微创手术占比", "出院病历3个工作日归档率", "检验危急值处理率", "临床路径入径率",
}
SPECIAL_METRICS = {"DDDs", "医疗不良事件上报数"}
# 不参与排名（仅显示数值）的指标
NO_RANK_METRICS = {"死亡例数（住院）"}


def parse_val(s):
    if s is None:
        return None
    s = str(s).strip().replace(",", "").replace("%", "")
    if s == "" or s == "-":
        return None
    try:
        return float(s) if "." in s else int(s)
    except Exception:
        return None


def norm_key(name):
    return re.sub(r"[（(].*?[)）]|\s+", "", str(name))


def fmt_val(metric, v):
    if v is None:
        return ""
    if metric in COUNT_METRICS:
        return str(int(round(v)))
    if metric in PERCENT_METRICS:
        return f"{v:.2f}%"
    if metric == "CMI值":
        return f"{v:.2f}"
    if isinstance(v, float):
        return f"{v:.2f}" if v != int(v) else str(int(v))
    return str(v)


def fmt_change(metric, cur_val, prev_val):
    if cur_val is None or prev_val is None or prev_val == 0:
        return ('\u2014', CLR_GRAY)
    pct = (cur_val - prev_val) / prev_val * 100
    sign = '+' if pct >= 0 else ''
    label = f"{sign}{pct:.1f}%"
    if metric in SPECIAL_METRICS or metric in NO_RANK_METRICS:
        return (label, CLR_GRAY)
    is_good = (pct >= 0 and metric in POSITIVE) or (pct < 0 and metric in NEGATIVE)
    return (label, CLR_GREEN if is_good else CLR_RED)

# ---------- 上月数据（可选） ----------
PREV_CSV = Path(_a.prev) if _a.prev else None
PREV_DATA = {}
HAS_PREV = False
if PREV_CSV and PREV_CSV.exists():
    with open(PREV_CSV, encoding='utf-8-sig') as _fp:
        _prev_hdrs = csv.DictReader(_fp).fieldnames or []
        for _row in csv.DictReader(open(PREV_CSV, encoding='utf-8-sig')):
            _dept_alias = DEPT_ALIASES.get(_row['科室'], _row['科室'])
            PREV_DATA[norm_key(_dept_alias)] = {h: parse_val(_row.get(h, '')) for h in _prev_hdrs if h != '科室'}
    HAS_PREV = True

# ---------- 读取 ----------
ddds_targets = {}
with open(DDDS_TARGET_CSV, encoding="utf-8-sig") as f:
    for row in csv.DictReader(f):
        t = parse_val(row.get("DDD目标值", ""))
        if t is not None:
            ddds_targets[norm_key(row["科室"])] = t

with open(DATA_CSV, encoding="utf-8-sig") as f:
    reader = csv.DictReader(f)
    headers = reader.fieldnames
    data = []
    for row in reader:
        rec = {"科室": row["科室"]}
        # 读取 DDDs达标 列（原始值：达标/不达标）
        rec["_ddds_dadab_raw"] = row.get("DDDs达标", "").strip() if row.get("DDDs达标") else None
        for m in headers[1:]:
            if m == "DDDs达标":
                continue
            rec[m] = parse_val(row.get(m))
        data.append(rec)

# 排名队列：有出院人次的科室（核心临床科室）
rank_depts = [r for r in data if r["出院人次"] is not None]
print(f"总科室 {len(data)}，参与排名的核心临床科室 {len(rank_depts)}")


def dadab(rec, metric):
    if metric == "DDDs":
        # 直接使用 CSV 中的 DDDs达标 列，不自行判断
        raw = rec.get("_ddds_dadab_raw")
        if raw == "达标":
            return True
        elif raw == "不达标":
            return False
        return None
    if metric == "医疗不良事件上报数":
        v = rec["医疗不良事件上报数"]
        dis = rec["出院人次"]
        if v is None or dis is None or dis == 0:
            return None
        return v >= dis * 0.025
    return None


# ---------- 计算排名与 ABCD（仅对 rank_depts）----------
for rec in rank_depts:
    rec["_rank"] = {}

for metric in ALL_METRICS:
    if metric in SPECIAL_METRICS:
        for rec in rank_depts:
            d = dadab(rec, metric)
            rec["_rank"][metric] = {"rank": None, "total": 0, "value": rec[metric],
                                    "quartile": ("A" if d else "D") if d is not None else None, "dadab": d}
        continue
    if metric in NO_RANK_METRICS:
        for rec in rank_depts:
            rec["_rank"][metric] = {"rank": None, "total": 0, "value": rec[metric], "quartile": None}
        continue
    vals = [(r["科室"], r[metric]) for r in rank_depts if r[metric] is not None]
    if not vals:
        for rec in rank_depts:
            rec["_rank"][metric] = {"rank": None, "total": 0, "value": rec[metric], "quartile": None}
        continue
    sv = sorted(vals, key=lambda x: x[1], reverse=metric in POSITIVE)
    ranks = {}
    idx = 0
    cur = 1
    while idx < len(sv):
        grp = [sv[idx]]
        j = idx + 1
        while j < len(sv) and sv[j][1] == sv[idx][1]:
            grp.append(sv[j]); j += 1
        for dpt in grp:
            ranks[dpt[0]] = cur
        cur += len(grp)
        idx = j
    total = len(sv)
    max_rank = max(ranks.values()) if ranks else 0
    all_same = len(set(v for _, v in sv)) == 1
    for rec in rank_depts:
        r = ranks.get(rec["科室"])
        if r is None or total == 0:
            q = None
        elif all_same:
            q = "D"
        elif r == max_rank:
            q = "D"
        elif r <= total // 4:
            q = "A"
        elif r <= total // 2:
            q = "B"
        elif r <= total * 3 // 4:
            q = "C"
        else:
            q = "D"
        rec["_rank"][metric] = {"rank": r, "total": total, "value": rec[metric], "quartile": q}
        if metric == "临床路径入径率" and rec[metric] is not None:
            rec["_rank"][metric]["quartile"] = "A" if rec[metric] > 50 else "D"
        if metric == "检验危急值处理率" and rec[metric] is not None and rec[metric] == 100.0:
            rec["_rank"][metric]["quartile"] = "A"


# ---------- 雷达图 ----------
def draw_radar(dept_rec, group_name, metrics, out_path):
    filt = [m for m in metrics if dept_rec[m] is not None]
    if len(filt) < 3:
        return False
    N = len(filt)
    angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()
    angles += angles[:1]
    q2r = {'D': 0.25, 'C': 0.5, 'B': 0.75, 'A': 1.0, None: 0.5}
    radii = [q2r.get(dept_rec["_rank"][m]["quartile"], 0.5) for m in filt]
    radii_c = radii + radii[:1]
    fig, ax = plt.subplots(figsize=(5.0, 5.0), subplot_kw=dict(polar=True))
    fig.patch.set_facecolor('white')
    for r, lab, col in zip([0.25, 0.5, 0.75, 1.0], ['D', 'C', 'B', 'A'],
                           [QUARTILE_RADAR_COLOR['D'], QUARTILE_RADAR_COLOR['C'],
                            QUARTILE_RADAR_COLOR['B'], QUARTILE_RADAR_COLOR['A']]):
        ax.add_artist(plt.Circle((0, 0), r, transform=ax.transData._b, fill=True,
                                 facecolor=col, alpha=0.06, edgecolor='none', zorder=0))
        ax.add_artist(plt.Circle((0, 0), r, transform=ax.transData._b, fill=False,
                                 edgecolor='#CCCCCC', linewidth=0.8, zorder=1))
    for angle in angles[:-1]:
        ax.plot([angle, angle], [0, 1], color='#CCCCCC', linewidth=0.5, zorder=1)
    ax.fill(angles, radii_c, color="#1F75FE", alpha=0.18, zorder=2)
    ax.plot(angles, radii_c, color="#1F3A5F", linewidth=2, marker='o', markersize=4,
            markerfacecolor='white', markeredgewidth=1.5, zorder=3)
    for angle, radius, m in zip(angles[:-1], radii, filt):
        qi = dept_rec["_rank"][m]["quartile"]
        if qi in QUARTILE_RADAR_COLOR:
            ax.text(angle, min(radius + 0.08, 1.05), qi, ha='center', va='center', fontsize=9,
                    fontweight='bold', color=QUARTILE_RADAR_COLOR[qi],
                    bbox=dict(boxstyle='circle,pad=0.15', facecolor='white',
                              edgecolor=QUARTILE_RADAR_COLOR[qi], linewidth=0.8), zorder=5)
        ax.text(angle, 1.25, m, ha='center', va='center', fontsize=7, fontproperties=_SYSTEM_FONT,
                color='#333333', bbox=dict(boxstyle='round,pad=0.15', facecolor='#FFFFF0',
                                           edgecolor='#CCCCCC', linewidth=0.5), zorder=5)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels([''] * len(filt))
    ax.set_yticks([0.25, 0.5, 0.75, 1.0])
    ax.set_yticklabels(['D', 'C', 'B', 'A'], fontsize=8, fontweight='bold', fontproperties=_SYSTEM_FONT)
    for tl, col in zip(ax.get_yticklabels(),
                       [QUARTILE_RADAR_COLOR['D'], QUARTILE_RADAR_COLOR['C'],
                        QUARTILE_RADAR_COLOR['B'], QUARTILE_RADAR_COLOR['A']]):
        tl.set_color(col)
    ax.set_ylim(0, 1.6)
    ax.set_title(f"{dept_rec['科室']} — {group_name}", fontsize=11, fontweight='bold',
                 pad=20, color='#1F3A5F', fontproperties=_SYSTEM_FONT)
    plt.tight_layout(pad=1.0)
    fig.savefig(out_path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return True


# ---------- Word ----------
def set_cell_font(cell, size=9, bold=False, color=CLR_BLACK, align=WD_ALIGN_PARAGRAPH.CENTER):
    for p in cell.paragraphs:
        p.alignment = align
        for run in p.runs:
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = color
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def shade(cell, color):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), color)
    sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)


def set_col_width(cell, twips):
    tc = cell._tc.get_or_add_tcPr()
    for child in list(tc):
        if child.tag.endswith('tcW'):
            tc.remove(child)
    w = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(twips))
    w.set(qn('w:type'), 'dxa')
    tc.append(w)


def make_report(rec):
    dept = rec["科室"]
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = '微软雅黑'
    st.font.size = Pt(10)
    st.font.color.rgb = CLR_BLACK
    st._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    for sec in doc.sections:
        sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)

    t = doc.add_heading("", level=0)
    run = t.add_run("科室质量控制数据报表")
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(20); run.font.color.rgb = CLR_TITLE
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(f"{dept}（2026年7月）")
    run.font.size = Pt(14); run.font.bold = True; run.font.color.rgb = CLR_TITLE
    run.font.name = "微软雅黑"; run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    expl = doc.add_paragraph(); expl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = expl.add_run("说明：评级按全院核心临床科室排名分四档，A（前25%，绿）、B（25%-50%，蓝）、C（50%-75%，橙）、D（后25%，红）；DDDs、医疗不良事件上报数为达标/不达标。")
    run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.font.name = "微软雅黑"; run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    prev_dept = PREV_DATA.get(norm_key(dept), {})

    # 数据表（5列：指标名称 | 本月 | 上月 | 增减 | 评级）
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Normal Table'
    tblPr = table._tbl.tblPr
    for child in list(tblPr):
        if child.tag.endswith('tblW'):
            tblPr.remove(child)
    tw = OxmlElement('w:tblW'); tw.set(qn('w:w'), '12000'); tw.set(qn('w:type'), 'dxa'); tblPr.append(tw)
    borders = OxmlElement('w:tblBorders')
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{e}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0'); b.set(qn('w:color'), 'AAAAAA')
        borders.append(b)
    tblPr.append(borders)

    hdr = table.rows[0].cells
    for i, txt in enumerate(["指标名称", "本月", "上月", "增减", "评级"]):
        hdr[i].text = txt
        set_cell_font(hdr[i], size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(hdr[i], CLR_HEADER_BG)
    set_col_width(hdr[0], 4200); set_col_width(hdr[1], 1700)
    set_col_width(hdr[2], 1700); set_col_width(hdr[3], 1800); set_col_width(hdr[4], 1400)

    for metric in ALL_METRICS:
        ri = rec["_rank"].get(metric, {})
        v = rec[metric]
        q = ri.get("quartile")
        dad = ri.get("dadab")
        prev_v = prev_dept.get(metric)
        cells = table.add_row().cells
        cells[0].text = metric
        set_cell_font(cells[0], size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
        cells[1].text = fmt_val(metric, v)
        set_cell_font(cells[1], size=9)
        cells[2].text = fmt_val(metric, prev_v)
        set_cell_font(cells[2], size=9, color=CLR_GRAY)
        chg_label, chg_color = fmt_change(metric, v, prev_v)
        cells[3].text = chg_label
        set_cell_font(cells[3], size=9, color=chg_color)
        set_col_width(cells[0], 4200); set_col_width(cells[1], 1700)
        set_col_width(cells[2], 1700); set_col_width(cells[3], 1800); set_col_width(cells[4], 1400)
        if metric in SPECIAL_METRICS and dad is not None:
            label = "达标" if dad else "不达标"
            p = cells[4].paragraphs[0]
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            nr = p.add_run(label)
            nr.font.size = Pt(9); nr.font.bold = True
            nr.font.color.rgb = CLR_GREEN if dad else CLR_RED
            nr.font.name = "微软雅黑"
        elif q:
            p = cells[4].paragraphs[0]
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            nr = p.add_run(q)
            nr.font.size = Pt(9); nr.font.bold = True
            nr.font.color.rgb = QUARTILE_COLOR[q]
            nr.font.name = "微软雅黑"

    # 雷达图（两个并排）
    doc.add_paragraph()
    rt = doc.add_table(rows=1, cols=2)
    rt.alignment = WD_TABLE_ALIGNMENT.CENTER
    rt.style = 'Normal Table'
    rp1 = RADAR_DIR / f"{dept}_能力及效率.png"
    rp2 = RADAR_DIR / f"{dept}_医疗质量与安全.png"
    ok1 = draw_radar(rec, "能力及效率", GROUPS["能力及效率"], rp1)
    ok2 = draw_radar(rec, "医疗质量与安全", GROUPS["医疗质量与安全"], rp2)
    c00 = rt.cell(0, 0); c01 = rt.cell(0, 1)
    if ok1:
        p = c00.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(rp1), width=Inches(2.6))
    else:
        c00.text = "（无能力及效率数据）"
        set_cell_font(c00, size=8, color=RGBColor(0x99, 0x99, 0x99))
    if ok2:
        p = c01.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(rp2), width=Inches(2.6))
    else:
        c01.text = "（无医疗质量与安全数据）"
        set_cell_font(c01, size=8, color=RGBColor(0x99, 0x99, 0x99))

    out = OUT_DIR / f"{dept}质控报表_7月.docx"
    doc.save(str(out))
    return out


# ---------- 生成 + 打包 ----------
generated = []
for rec in rank_depts:
    generated.append(make_report(rec))

print(f"已生成 {len(generated)} 个科室报表")
with zipfile.ZipFile(str(ZIP_PATH), 'w', zipfile.ZIP_DEFLATED) as zf:
    for p in generated:
        zf.write(str(p), arcname=p.name)
print(f"ZIP: {ZIP_PATH} ({ZIP_PATH.stat().st_size/1024/1024:.1f} MB)")
print("DONE")